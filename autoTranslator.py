# -*- coding: utf-8 -*-
"""
Created on Tue Nov  1 16:23:20 2022

@author: Noriko
"""

import deepl
import docx


# 翻译一个字符串，返回翻译过的字符串内容
def translate(sentence: str, target: str, source: str=None):
    # 身份认证
    auth_key = "personal info"
    translator = deepl.Translator(auth_key) 
    # 调用deepl翻译
    tdSentence = translator.translate_text(sentence, target_lang=target, source_lang=source)
    return tdSentence

# 翻译文档
def translateDocx(filePath: str, target: str, source: str=None):
    file=docx.Document(filePath)
    doc = docx.Document()
    # 读取每一段 进行翻译
    for para in file.paragraphs:
        doc.add_paragraph(str(translate(para.text.replace("「","\"").replace("」","\""), target, source)))
    # 储存到新docx
    doc.save("td_"+filePath)

    
if __name__ == "__main__":
    filePath = "test.docx"
    translateDocx(filePath, "ZH", "JA")
    
